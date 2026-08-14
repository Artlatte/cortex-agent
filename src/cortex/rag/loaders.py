"""Document loaders for the supported file formats.

Each loader turns a single file into a single :class:`Document`. Optional
third-party parsers (pypdf, python-docx) are imported lazily so that a missing
dependency surfaces as a :class:`DocumentLoadError` with an install hint rather
than crashing the import of the whole RAG module.
"""

from __future__ import annotations

from pathlib import Path
from typing import Protocol, runtime_checkable

from cortex.errors import DocumentLoadError
from cortex.rag.document import Document

SUPPORTED_SUFFIXES = frozenset({".txt", ".md", ".markdown", ".pdf", ".docx"})
_TEXT_SUFFIXES = frozenset({".txt", ".md", ".markdown"})


@runtime_checkable
class DocumentLoader(Protocol):
    """Structural interface for anything that loads documents from a path."""

    def load(self, path: str | Path) -> list[Document]: ...


class TextLoader:
    """Loads UTF-8 plain-text files (``.txt``, ``.md``, ``.markdown``)."""

    def load(self, path: str | Path) -> list[Document]:
        path = Path(path)
        try:
            text = path.read_text(encoding="utf-8")
        except OSError as exc:
            raise DocumentLoadError(f"failed to read {path}: {exc}") from exc
        metadata = {
            "source": str(path),
            "filename": path.name,
            "filetype": path.suffix.lstrip(".").lower(),
        }
        return [Document(page_content=text, metadata=metadata)]


class PdfLoader:
    """Loads PDF files with :mod:`pypdf`, one document per file."""

    def load(self, path: str | Path) -> list[Document]:
        path = Path(path)
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise DocumentLoadError(
                "pypdf is required to load PDF files; install it with `pip install pypdf`"
            ) from exc
        try:
            reader = PdfReader(str(path))
            if not reader.pages:
                raise DocumentLoadError(f"failed to parse PDF {path}: no pages found")
            pages = [page.extract_text() or "" for page in reader.pages]
        except DocumentLoadError:
            raise
        except Exception as exc:
            raise DocumentLoadError(f"failed to parse PDF {path}: {exc}") from exc
        metadata = {
            "source": str(path),
            "filename": path.name,
            "filetype": "pdf",
            "page_count": len(reader.pages),
        }
        return [Document(page_content="\n\n".join(pages), metadata=metadata)]


class DocxLoader:
    """Loads DOCX files with :mod:`python-docx` (paragraphs and table cells)."""

    def load(self, path: str | Path) -> list[Document]:
        path = Path(path)
        try:
            import docx
        except ImportError as exc:
            raise DocumentLoadError(
                "python-docx is required to load DOCX files; install it with "
                "`pip install python-docx`"
            ) from exc
        try:
            document = docx.Document(str(path))
        except Exception as exc:
            raise DocumentLoadError(f"failed to parse DOCX {path}: {exc}") from exc

        parts: list[str] = [p.text for p in document.paragraphs if p.text]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        metadata = {
            "source": str(path),
            "filename": path.name,
            "filetype": "docx",
            "paragraph_count": len(document.paragraphs),
        }
        return [Document(page_content="\n".join(parts), metadata=metadata)]


def get_loader(path: str | Path) -> DocumentLoader:
    """Return the loader matching ``path``'s suffix or raise on unknown types."""
    suffix = Path(path).suffix.lower()
    if suffix in _TEXT_SUFFIXES:
        return TextLoader()
    if suffix == ".pdf":
        return PdfLoader()
    if suffix == ".docx":
        return DocxLoader()
    raise DocumentLoadError(f"unsupported file type: {suffix}")
