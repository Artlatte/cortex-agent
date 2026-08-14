"""Tests for the document loaders and suffix dispatch."""

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from cortex.errors import DocumentLoadError
from cortex.rag import DocumentLoader, get_loader


def test_text_loader_txt_and_md(tmp_path) -> None:
    txt = tmp_path / "a.txt"
    txt.write_text("hello world", encoding="utf-8")
    loader = get_loader(txt)
    assert isinstance(loader, DocumentLoader)
    docs = loader.load(txt)
    assert len(docs) == 1
    assert docs[0].page_content == "hello world"
    assert docs[0].metadata["source"] == str(txt)
    assert docs[0].metadata["filename"] == "a.txt"
    assert docs[0].metadata["filetype"] == "txt"

    md = tmp_path / "b.md"
    md.write_text("# title\nbody", encoding="utf-8")
    docs = get_loader(md).load(md)
    assert docs[0].metadata["filetype"] == "md"
    assert docs[0].page_content == "# title\nbody"


def test_pdf_loader(tmp_path) -> None:
    pdf_path = tmp_path / "doc.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with pdf_path.open("wb") as fh:
        writer.write(fh)
    docs = get_loader(pdf_path).load(pdf_path)
    assert len(docs) == 1
    assert docs[0].metadata["page_count"] == 1
    assert docs[0].metadata["filetype"] == "pdf"


def test_docx_loader(tmp_path) -> None:
    docx_path = tmp_path / "doc.docx"
    document = DocxDocument()
    document.add_paragraph("Hello")
    document.add_paragraph("World")
    document.save(str(docx_path))
    docs = get_loader(docx_path).load(docx_path)
    assert len(docs) == 1
    assert "Hello" in docs[0].page_content
    assert "World" in docs[0].page_content
    assert docs[0].metadata["filetype"] == "docx"
    assert docs[0].metadata["paragraph_count"] >= 2


def test_garbage_pdf_raises_document_load_error(tmp_path) -> None:
    bad = tmp_path / "bad.pdf"
    bad.write_bytes(b"\x00\x01\x02this is definitely not a pdf")
    with pytest.raises(DocumentLoadError):
        get_loader(bad).load(bad)


def test_unknown_suffix_raises_document_load_error(tmp_path) -> None:
    unknown = tmp_path / "file.xyz"
    unknown.write_text("data", encoding="utf-8")
    with pytest.raises(DocumentLoadError):
        get_loader(unknown)
